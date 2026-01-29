#!/usr/bin/env python3
"""
Convert Markdown documentation to Microsoft Word format
Requires: pip install python-docx
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def convert_markdown_to_word(md_file_path, output_path):
    """Convert Markdown file to Word document"""
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    h1_font = h1_style.font
    h1_font.name = 'Arial'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    h2_font = h2_style.font
    h2_font.name = 'Arial'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(6)
    h2_style.paragraph_format.space_after = Pt(3)
    
    # Code style
    code_style = styles.add_style('CodeStyle', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle main title
        if line.startswith('# Intern Management System'):
            p = doc.add_paragraph(line, style='CustomTitle')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue
            
        # Handle headings
        if line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='CustomHeading1')
            i += 1
            continue
            
        if line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='CustomHeading2')
            i += 1
            continue
            
        if line.startswith('#### '):
            p = doc.add_paragraph(line[5:], style='CustomHeading2')
            p.style.font.bold = True
            p.style.font.size = Pt(11)
            i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            # Find end of code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                p = doc.add_paragraph()
                p.style = code_style
                for code_line in code_lines:
                    if code_line.strip():
                        run = p.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
            
            i += 1
            continue
        
        # Handle lists
        if line.startswith('- ') or line.startswith('* '):
            # Handle list items
            while i < len(lines) and (lines[i].startswith('- ') or lines[i].startswith('* ')):
                list_line = lines[i][2:].strip()
                p = doc.add_paragraph(list_line, style='List Bullet')
                i += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\. ', line):
            while i < len(lines) and re.match(r'^\d+\. ', lines[i]):
                list_line = re.sub(r'^\d+\. ', '', lines[i]).strip()
                p = doc.add_paragraph(list_line, style='List Number')
                i += 1
            continue
        
        # Handle inline code
        if '`' in line:
            parts = line.split('`')
            p = doc.add_paragraph()
            for j, part in enumerate(parts):
                if j % 2 == 0:
                    # Regular text
                    if part:
                        run = p.add_run(part)
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                else:
                    # Code text
                    run = p.add_run(part)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.style.font.name = 'Arial'
            p.style.font.size = Pt(11)
        
        i += 1
    
    # Add table of contents placeholder
    doc.add_page_break()
    p = doc.add_paragraph("Table of Contents", style='CustomHeading1')
    p.add_run("\n\n(Word's built-in Table of Contents can be inserted here)")
    
    # Save document
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

def main():
    """Main function"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(current_dir, 'PROJECT_DOCUMENTATION.md')
    output_file = os.path.join(current_dir, 'Intern_Management_System_Documentation.docx')
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found!")
        return
    
    try:
        convert_markdown_to_word(md_file, output_file)
        print(f"✅ Documentation successfully converted to Word format!")
        print(f"📄 Output file: {output_file}")
        print(f"📊 File size: {os.path.getsize(output_file) / 1024:.1f} KB")
    except Exception as e:
        print(f"❌ Error converting to Word: {e}")
        print("\n💡 To install required package, run:")
        print("   pip install python-docx")

if __name__ == "__main__":
    main()
